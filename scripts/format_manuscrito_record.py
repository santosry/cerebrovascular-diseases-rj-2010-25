"""Aplica estilo editorial reproduzível ao DOCX gerado pelo R Markdown."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PATH = ROOT / "results" / "manuscrito_RECORD.docx"

BLUE = "1F4E79"
DARK_BLUE = "17365D"
TEXT = "202124"
MUTED = "5F6368"
LIGHT = "F2F4F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


doc = Document(PATH)

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "MORBIMORTALIDADE CEREBROVASCULAR · RJ"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Relato STROBE–RECORD  ·  ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(TEXT)
normal.paragraph_format.line_spacing = 1.10
normal.paragraph_format.space_after = Pt(6)

for name, size, color, before, after in (
    ("Title", 22, BLUE, 30, 18),
    ("Subtitle", 12, MUTED, 6, 14),
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = name != "Subtitle"
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

title_seen = False
on_cover = False
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if paragraph.style.name == "Title":
        title_seen = True
        on_cover = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(48)
        paragraph.paragraph_format.space_after = Pt(24)
        paragraph.paragraph_format.keep_with_next = True
    elif on_cover and text != "17 de julho de 2026":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(14)
        for run in paragraph.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor.from_string(MUTED)
    elif text == "17 de julho de 2026":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(28)
        for run in paragraph.runs:
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(BLUE)
        paragraph.add_run().add_break(WD_BREAK.PAGE)
        on_cover = False
    if paragraph.style.name.startswith("Heading"):
        keep_with_next(paragraph)
    if text.startswith("Apêndice B") or text.startswith("Apêndice C"):
        paragraph.paragraph_format.page_break_before = True

for table in doc.tables:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    if table.rows:
        set_repeat_table_header(table.rows[0])
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    header_key = tuple(headers)
    widths = None
    if headers[:2] == ["sistema", "etapa"]:
        widths = [0.65, 0.45, 3.05, 0.75, 1.55]
    elif headers == ["ano", "internacoes", "obitos_hospitalares", "mortalidade_hospitalar_pct"]:
        widths = [0.65, 1.20, 1.65, 2.90]
    elif headers == ["ano", "permanencia_media", "permanencia_mediana", "custo_total", "custo_medio"]:
        widths = [0.60, 1.35, 1.35, 1.55, 1.55]
    elif headers and headers[0] == "vies":
        widths = [0.75, 1.75, 1.25, 1.15, 1.60]
    elif headers == ["item", "status", "evidencia"]:
        widths = [0.60, 1.10, 4.70]
    elif headers == ["referencia", "requisito", "status", "evidencia"]:
        widths = [0.75, 1.70, 0.90, 3.05]
    elif headers and headers[0] == "serie":
        widths = [2.25, 0.90, 0.45, 0.55, 0.70, 0.80, 0.65]
    elif headers and headers[0] == "sistema" and "p_bonferroni_entre_anos" in headers:
        widths = [0.60, 0.55, 0.45, 0.55, 0.85, 0.85, 0.75, 1.80]
    elif headers and headers[0] == "variavel":
        widths = [1.15, 0.70, 2.20, 0.75, 0.95, 0.65]
    if widths and len(widths) == len(table.columns):
        table.autofit = False
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                set_cell_width(cell, width)
    wide = len(table.columns) >= 5
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(7.5 if wide else 8.5)
                    if row_index == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(WHITE)

doc.core_properties.title = "Morbimortalidade cerebrovascular no Rio de Janeiro — relato STROBE–RECORD"
doc.core_properties.subject = "SIH/SUS 2010–2025 e SIM 2010–2024"
doc.core_properties.keywords = "RECORD, STROBE, SIH/SUS, SIM, cerebrovascular, Rio de Janeiro"
doc.save(PATH)
